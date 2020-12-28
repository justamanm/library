# pip install python-docx
# python-docx只能读取docx文件，不能读取doc文件
from docx import Document

# pip install pywin32
# win32可以操作doc文件
from win32com import client as wc

import os


class Word:
    def __init__(self):
        self.word = wc.Dispatch('Word.Application')
        # 或者使用下面的方法，使用启动独立的进程：
        # word = wc.DispatchEx('Word.Application')

    def doc2docx(self):
        path = os.getcwd()
        doc = self.word.Documents.Open(path + "/data.doc")
        doc.SaveAs(path + "/data.docx", 16)
        doc.Close()
        self.word.Quit()

    # 获取word表格中的内容
    def read_table(self):
        # 读取文档
        doc = Document("data.docx")  # filename为word文档
        # 获取文档中的表格，返回列表
        tables = doc.tables
        # print(len(tables))

        # 读取第1个表格
        tb1 = doc.tables[0]
        # 获取第一个表格的所有行，行下标从0开始
        rows = tb1.rows  # 获取表格的行数len(tb1.rows)
        # print(len(rows))

        # 读取表格的第一行的所有单元格
        # row_cells = tb1.rows[0].cells

        for i, row in enumerate(rows):
            for cell in row.cells:
                print(cell.text)
